#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
DocxEditor - LLM 原子化 Word 编辑工具
=====================================

设计：纯执行层，判断逻辑由 LLM 完成
索引：三轨制 - 段落 index / 表格 table_index / 图片 image_index，互不干扰

段落：
    get_outline()           获取大纲
    read_content(indices)   读取段落

表格：
    get_tables_outline()    获取表格概览
    read_table(table_index) 读取表格内容

图片：
    get_images_outline()    获取图片概览

通用：
    batch_update(ops)       批量修改（自动倒序）
    save(path)              保存

示例：
    from docx_editor import DocxEditor
    editor = DocxEditor('input.docx')

    # 段落
    editor.batch_update([{'op': 'delete', 'index': 50}])

    # 表格
    editor.batch_update([{'op': 'update_table_cell', 'table_index': 0, 'row': 1, 'col': 2, 'text': '新'}])

    # 图片
    editor.batch_update([{'op': 'resize_image', 'image_index': 0, 'width': 10.0}])

    # 引用刷新（让Word打开时自动更新目录/页码/交叉引用）
    editor.batch_update([{'op': 'update_fields_on_open'}])

    editor.save('output.docx')

详见 README_DOCX_EDITOR.md
"""

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.shape import WD_INLINE_SHAPE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import os
from typing import Optional, List, Dict, Union


class DocxEditor:
    """Word 文档编辑器 - 纯执行层"""

    def __init__(self, file_path: str):
        self.file_path = file_path
        self.doc = Document(file_path)
        self._refresh()

    def _refresh(self):
        """刷新段落、表格、图片索引"""
        self._paragraphs = list(self.doc.paragraphs)
        self._tables = list(self.doc.tables)
        self._images = list(self.doc.inline_shapes)

    # ==================== 查询层（Read）====================

    def get_outline(self) -> Dict:
        """
        获取文档大纲（极低Token消耗）

        仅返回标题结构，用于建立文档地图。

        Returns:
            {
                'total': 段落总数,
                'headings': [
                    {'index': 0, 'level': 1, 'text': '第一章 绪论'},
                    {'index': 15, 'level': 2, 'text': '1.1 研究背景'},
                    ...
                ]
            }
        """
        headings = []
        for i, para in enumerate(self._paragraphs):
            level = self._get_heading_level(para)
            if level and para.text.strip():
                headings.append({
                    'index': i,
                    'level': level,
                    'text': para.text.strip()[:100]
                })

        return {
            'total': len(self._paragraphs),
            'headings': headings
        }

    def read_content(self, indices: Union[int, List[int], range, str]) -> List[Dict]:
        """
        按需读取指定段落详情

        Args:
            indices: 可以是:
                - int: 单个索引
                - List[int]: 索引列表
                - range: 索引范围 range(10, 20)
                - str: 章节标题（读取该章节下所有内容）

        Returns:
            [
                {
                    'index': 10,
                    'text': '完整段落文本',
                    'style': 'Normal',
                    'is_heading': False,
                    'heading_level': None,
                    'is_empty': False,
                    'runs': [
                        {'text': '部分文本', 'bold': True, 'italic': False, 'font_size': 12}
                    ],
                    'format': {
                        'alignment': 'justify',
                        'line_spacing': 1.5,
                        'first_line_indent': 0.74,
                        'left_indent': 0
                    },
                    'xml': {
                        'has_numPr': True,  # 是否有自动编号
                        'style_name': 'List Paragraph'
                    }
                }
            ]
        """
        # 解析 indices
        if isinstance(indices, int):
            idx_list = [indices]
        elif isinstance(indices, str):
            # 按章节标题查找
            idx_list = self._get_section_indices(indices)
        elif isinstance(indices, range):
            idx_list = list(indices)
        else:
            idx_list = list(indices)

        results = []
        for i in idx_list:
            if not 0 <= i < len(self._paragraphs):
                continue

            para = self._paragraphs[i]
            pf = para.paragraph_format
            level = self._get_heading_level(para)

            # 获取 runs 详情
            runs = []
            for r in para.runs:
                run_info = {
                    'text': r.text,
                    'bold': r.bold,
                    'italic': r.italic,
                }
                if r.font.size:
                    run_info['font_size'] = r.font.size.pt
                runs.append(run_info)

            # 获取格式
            fmt = {
                'alignment': self._alignment_to_str(para.alignment),
                'line_spacing': pf.line_spacing,
            }
            if pf.first_line_indent:
                fmt['first_line_indent'] = round(pf.first_line_indent.cm, 2) if hasattr(pf.first_line_indent, 'cm') else None
            if pf.left_indent:
                fmt['left_indent'] = round(pf.left_indent.cm, 2) if hasattr(pf.left_indent, 'cm') else None

            # 获取 XML 级别信息
            xml_info = {
                'has_numPr': False,
                'style_name': para.style.name if para.style else None
            }
            pPr = para._element.pPr
            if pPr is not None:
                if pPr.find(qn('w:numPr')) is not None:
                    xml_info['has_numPr'] = True

            # 检查段落是否包含嵌入对象（图片、OLE对象等）
            # 这些内容不会体现在 text 属性中，但删除段落会一并删除
            drawings = para._element.findall('.//' + qn('w:drawing'))
            objects = para._element.findall('.//' + qn('w:object'))
            has_embedded = len(drawings) > 0 or len(objects) > 0

            # is_truly_empty: 真正为空（无文字且无嵌入对象），可安全删除
            truly_empty = self._is_truly_empty(para)

            results.append({
                'index': i,
                'text': para.text,
                'style': para.style.name if para.style else 'Normal',
                'is_heading': level is not None,
                'heading_level': level,
                'is_empty': not para.text.strip(),  # 文本为空（可能包含图片）
                'is_truly_empty': truly_empty,       # 真正为空（可安全删除）
                'has_embedded': has_embedded,        # 是否包含图片/OLE等嵌入对象
                'runs': runs,
                'format': fmt,
                'xml': xml_info
            })

        return results

    # ==================== 表格查询层 ====================

    def get_tables_outline(self) -> List[Dict]:
        """
        获取所有表格概览

        Returns:
            [
                {
                    'table_index': 0,
                    'rows': 5,
                    'cols': 3,
                    'preview': '第一行第一列文本...'
                }
            ]
        """
        results = []
        for i, table in enumerate(self._tables):
            rows = len(table.rows)
            cols = len(table.columns) if table.rows else 0

            # 获取预览（第一个单元格的文本）
            preview = ''
            if rows > 0 and cols > 0:
                try:
                    cell_text = table.cell(0, 0).text.strip()
                    preview = cell_text[:50] + ('...' if len(cell_text) > 50 else '')
                except:
                    pass

            results.append({
                'table_index': i,
                'rows': rows,
                'cols': cols,
                'preview': preview
            })

        return results

    def read_table(self, table_index: int) -> Dict:
        """
        读取指定表格的完整内容

        Args:
            table_index: 表格索引

        Returns:
            {
                'table_index': 0,
                'rows': 5,
                'cols': 3,
                'data': [
                    ['单元格1', '单元格2', '单元格3'],
                    ['第二行1', '第二行2', '第二行3'],
                    ...
                ],
                'merged_cells': [(0,0,1,2), ...]  # 合并单元格信息 (row_start, col_start, row_end, col_end)
            }
        """
        if not 0 <= table_index < len(self._tables):
            raise IndexError(f"表格索引超出范围: {table_index}, 共 {len(self._tables)} 个表格")

        table = self._tables[table_index]
        rows = len(table.rows)
        cols = len(table.columns) if table.rows else 0

        # 读取所有单元格内容
        data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                # 单元格可能包含多个段落，用换行连接
                cell_text = '\n'.join(p.text for p in cell.paragraphs)
                row_data.append(cell_text)
            data.append(row_data)

        return {
            'table_index': table_index,
            'rows': rows,
            'cols': cols,
            'data': data
        }

    # ==================== 图片查询层 ====================

    def get_images_outline(self) -> List[Dict]:
        """
        获取所有图片概览

        Returns:
            [
                {
                    'image_index': 0,
                    'type': 'picture',  # picture/chart/smart_art/other
                    'width_cm': 5.2,
                    'height_cm': 3.1,
                    'desc': '图片描述(如有)'
                }
            ]
        """
        results = []
        for i, shape in enumerate(self._images):
            # 判断类型
            shape_type = 'other'
            if shape.type == WD_INLINE_SHAPE_TYPE.PICTURE:
                shape_type = 'picture'
            elif shape.type == WD_INLINE_SHAPE_TYPE.CHART:
                shape_type = 'chart'
            elif shape.type == WD_INLINE_SHAPE_TYPE.SMART_ART:
                shape_type = 'smart_art'
            elif shape.type == WD_INLINE_SHAPE_TYPE.LINKED_PICTURE:
                shape_type = 'linked_picture'

            # 获取尺寸（EMU转厘米）
            width_cm = round(shape.width / 914400 * 2.54, 2) if shape.width else None
            height_cm = round(shape.height / 914400 * 2.54, 2) if shape.height else None

            results.append({
                'image_index': i,
                'type': shape_type,
                'width_cm': width_cm,
                'height_cm': height_cm
            })

        return results

    # ==================== 修改层（Update）====================

    def batch_update(self, operations: List[Dict]) -> Dict:
        """
        统一修改入口（自动倒序执行防止索引漂移）

        Args:
            operations: 操作列表，每个操作是一个字典：

            删除段落（自动保护含图片的段落）：
                {'op': 'delete', 'index': 50}
                {'op': 'delete', 'index': 50, 'force': True}  # 强制删除（慎用）

            插入段落：
                {'op': 'insert', 'index': 10, 'position': 'after'|'before', 'text': '新内容', 'style': 'Normal'}

            修改样式：
                {'op': 'update_style', 'index': 20,
                 'style': 'Normal',  # 段落样式
                 'font': {'name': '宋体', 'size': 12, 'bold': False},
                 'alignment': 'justify',  # left/center/right/justify
                 'indent': {'first_line': 0.74, 'left': 0},  # 单位：厘米
                 'spacing': {'before': 0, 'after': 0, 'line': 1.5}
                }

            替换文本（支持正则）：
                {'op': 'replace_text', 'index': 30,
                 'pattern': r'^（\\d）\\s*',  # 正则模式
                 'replacement': '',  # 替换内容
                 'regex': True  # 是否正则模式，默认True
                }

            全局替换（所有段落）：
                {'op': 'replace_text_global',
                 'pattern': '旧文本',
                 'replacement': '新文本',
                 'regex': False
                }

            清理XML属性：
                {'op': 'clean_xml', 'index': 15,
                 'remove': ['numPr'],  # 要移除的XML元素
                 'style': 'Normal',  # 可选：同时设置样式
                 'indent': {'first_line': 0.74, 'left': 0}  # 可选：同时设置缩进
                }

            设置段落文本：
                {'op': 'set_text', 'index': 25, 'text': '新的段落内容'}

            ===== 表格操作（使用 table_index，与段落 index 互不干扰）=====

            修改单元格：
                {'op': 'update_table_cell', 'table_index': 0, 'row': 1, 'col': 2, 'text': '新内容'}

            替换单元格文本（支持正则）：
                {'op': 'replace_table_cell', 'table_index': 0, 'row': 1, 'col': 2,
                 'pattern': '旧文本', 'replacement': '新文本', 'regex': False}

            批量修改整行：
                {'op': 'update_table_row', 'table_index': 0, 'row': 1,
                 'texts': ['列1', '列2', '列3']}

            批量修改整列：
                {'op': 'update_table_col', 'table_index': 0, 'col': 1,
                 'texts': ['行1', '行2', '行3']}

            ===== 图片操作（使用 image_index）=====

            删除图片：
                {'op': 'delete_image', 'image_index': 0}

            调整图片大小：
                {'op': 'resize_image', 'image_index': 0, 'width': 10.0}  # 宽度cm，高度自动
                {'op': 'resize_image', 'image_index': 0, 'width': 10.0, 'height': 5.0}

            在段落中插入图片：
                {'op': 'insert_image', 'index': 10, 'path': '/path/to/img.png', 'width': 6.0}

            ===== 引用/目录刷新 =====

            强制Word打开时刷新所有引用（目录、页码、交叉引用）：
                {'op': 'update_fields_on_open'}

        Returns:
            {
                'success': 成功数,
                'failed': 失败数,
                'details': [{'op': '...', 'index': ..., 'status': 'ok'|'error', 'error': '...'}]
            }
        """
        # 按 index 倒序排列（防止索引漂移）
        # insert 和 delete 会改变后续索引，所以从后往前执行
        sorted_ops = sorted(
            operations,
            key=lambda x: x.get('index', float('inf')),
            reverse=True
        )

        results = {'success': 0, 'failed': 0, 'details': []}

        for op in sorted_ops:
            op_type = op.get('op', '')
            index = op.get('index')
            detail = {'op': op_type, 'index': index, 'status': 'ok'}

            try:
                if op_type == 'delete':
                    self._op_delete(index, force=op.get('force', False))

                elif op_type == 'insert':
                    new_idx = self._op_insert(
                        index,
                        op.get('text', ''),
                        op.get('position', 'after'),
                        op.get('style')
                    )
                    detail['new_index'] = new_idx

                elif op_type == 'update_style':
                    self._op_update_style(index, op)

                elif op_type == 'replace_text':
                    changed = self._op_replace_text(
                        index,
                        op.get('pattern', ''),
                        op.get('replacement', ''),
                        op.get('regex', True)
                    )
                    detail['changed'] = changed

                elif op_type == 'replace_text_global':
                    count = self._op_replace_text_global(
                        op.get('pattern', ''),
                        op.get('replacement', ''),
                        op.get('regex', False)
                    )
                    detail['replaced_count'] = count

                elif op_type == 'clean_xml':
                    self._op_clean_xml(index, op)

                elif op_type == 'set_text':
                    self._op_set_text(index, op.get('text', ''))

                # ===== 表格操作 =====
                elif op_type == 'update_table_cell':
                    self._op_update_table_cell(
                        op.get('table_index'),
                        op.get('row'),
                        op.get('col'),
                        op.get('text', '')
                    )
                    detail['table_index'] = op.get('table_index')
                    detail['row'] = op.get('row')
                    detail['col'] = op.get('col')

                elif op_type == 'replace_table_cell':
                    changed = self._op_replace_table_cell(
                        op.get('table_index'),
                        op.get('row'),
                        op.get('col'),
                        op.get('pattern', ''),
                        op.get('replacement', ''),
                        op.get('regex', False)
                    )
                    detail['table_index'] = op.get('table_index')
                    detail['changed'] = changed

                elif op_type == 'update_table_row':
                    self._op_update_table_row(
                        op.get('table_index'),
                        op.get('row'),
                        op.get('texts', [])
                    )
                    detail['table_index'] = op.get('table_index')
                    detail['row'] = op.get('row')

                elif op_type == 'update_table_col':
                    self._op_update_table_col(
                        op.get('table_index'),
                        op.get('col'),
                        op.get('texts', [])
                    )
                    detail['table_index'] = op.get('table_index')
                    detail['col'] = op.get('col')

                # ===== 图片操作 =====
                elif op_type == 'delete_image':
                    self._op_delete_image(op.get('image_index'))
                    detail['image_index'] = op.get('image_index')

                elif op_type == 'resize_image':
                    self._op_resize_image(
                        op.get('image_index'),
                        op.get('width'),
                        op.get('height')
                    )
                    detail['image_index'] = op.get('image_index')

                elif op_type == 'insert_image':
                    self._op_insert_image(
                        index,
                        op.get('path', ''),
                        op.get('width'),
                        op.get('height')
                    )

                # ===== 引用刷新 =====
                elif op_type == 'update_fields_on_open':
                    self._op_update_fields_on_open()

                else:
                    detail['status'] = 'error'
                    detail['error'] = f'未知操作类型: {op_type}'
                    results['failed'] += 1
                    results['details'].append(detail)
                    continue

                results['success'] += 1

            except Exception as e:
                detail['status'] = 'error'
                detail['error'] = str(e)
                results['failed'] += 1

            results['details'].append(detail)

        return results

    def save(self, path: Optional[str] = None):
        """保存文档"""
        self.doc.save(path or self.file_path)

    # ==================== 内部操作实现 ====================

    def _op_delete(self, index: int, force: bool = False):
        """
        删除段落

        Args:
            index: 段落索引
            force: 是否强制删除（即使包含图片等嵌入对象）
                   默认 False，会拒绝删除包含嵌入对象的段落
        """
        if not 0 <= index < len(self._paragraphs):
            raise IndexError(f"索引超出范围: {index}")
        para = self._paragraphs[index]

        # 安全检查：防止误删包含图片/OLE对象的段落
        if not force and not self._is_truly_empty(para):
            raise ValueError(
                f"段落 {index} 包含嵌入对象（图片/OLE等），不能作为空段落删除。"
                f"如确需删除，请使用 force=True"
            )

        para._element.getparent().remove(para._element)
        self._refresh()

    def _op_insert(self, index: int, text: str, position: str, style: Optional[str]) -> int:
        """插入段落"""
        if not 0 <= index < len(self._paragraphs):
            raise IndexError(f"索引超出范围: {index}")

        ref_para = self._paragraphs[index]
        new_p = OxmlElement('w:p')

        if position == 'before':
            ref_para._element.addprevious(new_p)
            new_idx = index
        else:
            ref_para._element.addnext(new_p)
            new_idx = index + 1

        self._refresh()
        new_para = self._paragraphs[new_idx]
        new_para.add_run(text)

        if style:
            try:
                new_para.style = style
            except:
                pass

        return new_idx

    def _op_update_style(self, index: int, op: Dict):
        """修改样式"""
        if not 0 <= index < len(self._paragraphs):
            raise IndexError(f"索引超出范围: {index}")

        para = self._paragraphs[index]

        # 段落样式
        if 'style' in op:
            try:
                para.style = op['style']
            except:
                pass

        # 对齐
        if 'alignment' in op:
            align_map = {
                'left': WD_ALIGN_PARAGRAPH.LEFT,
                'center': WD_ALIGN_PARAGRAPH.CENTER,
                'right': WD_ALIGN_PARAGRAPH.RIGHT,
                'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
            }
            para.alignment = align_map.get(op['alignment'].lower())

        # 缩进
        if 'indent' in op:
            ind = op['indent']
            pf = para.paragraph_format
            if 'first_line' in ind:
                pf.first_line_indent = Cm(ind['first_line'])
            if 'left' in ind:
                pf.left_indent = Cm(ind['left'])
            if 'right' in ind:
                pf.right_indent = Cm(ind['right'])

        # 间距
        if 'spacing' in op:
            sp = op['spacing']
            pf = para.paragraph_format
            if 'before' in sp:
                pf.space_before = Pt(sp['before'])
            if 'after' in sp:
                pf.space_after = Pt(sp['after'])
            if 'line' in sp:
                pf.line_spacing = sp['line']

        # 字体（应用到所有 runs）
        if 'font' in op:
            f = op['font']
            for run in para.runs:
                if 'name' in f:
                    run.font.name = f['name']
                    # 同时设置中文字体
                    if run._element.rPr is None:
                        run._element.get_or_add_rPr()
                    rFonts = run._element.rPr.get_or_add_rFonts()
                    rFonts.set(qn('w:eastAsia'), f['name'])
                if 'size' in f:
                    run.font.size = Pt(f['size'])
                if 'bold' in f:
                    run.bold = f['bold']
                if 'italic' in f:
                    run.italic = f['italic']

    def _op_replace_text(self, index: int, pattern: str, replacement: str, regex: bool) -> bool:
        """替换单段落文本"""
        if not 0 <= index < len(self._paragraphs):
            raise IndexError(f"索引超出范围: {index}")

        para = self._paragraphs[index]
        original = para.text

        if regex:
            new_text = re.sub(pattern, replacement, original)
        else:
            new_text = original.replace(pattern, replacement)

        if new_text != original:
            self._set_paragraph_text(para, new_text)
            return True
        return False

    def _op_replace_text_global(self, pattern: str, replacement: str, regex: bool) -> int:
        """全局替换文本"""
        count = 0
        for para in self._paragraphs:
            original = para.text
            if regex:
                new_text = re.sub(pattern, replacement, original)
            else:
                new_text = original.replace(pattern, replacement)

            if new_text != original:
                self._set_paragraph_text(para, new_text)
                count += 1
        return count

    def _op_clean_xml(self, index: int, op: Dict):
        """清理XML属性"""
        if not 0 <= index < len(self._paragraphs):
            raise IndexError(f"索引超出范围: {index}")

        para = self._paragraphs[index]
        remove_list = op.get('remove', [])

        pPr = para._element.pPr
        if pPr is not None:
            for tag in remove_list:
                elem = pPr.find(qn(f'w:{tag}'))
                if elem is not None:
                    pPr.remove(elem)

        # 可选：同时设置样式
        if 'style' in op:
            try:
                para.style = op['style']
            except:
                pass

        # 可选：同时设置缩进
        if 'indent' in op:
            ind = op['indent']
            pf = para.paragraph_format
            if 'first_line' in ind:
                pf.first_line_indent = Cm(ind['first_line'])
            if 'left' in ind:
                pf.left_indent = Cm(ind['left'])

    def _op_set_text(self, index: int, text: str):
        """设置段落文本"""
        if not 0 <= index < len(self._paragraphs):
            raise IndexError(f"索引超出范围: {index}")

        para = self._paragraphs[index]
        self._set_paragraph_text(para, text)

    # ==================== 表格操作实现 ====================

    def _op_update_table_cell(self, table_index: int, row: int, col: int, text: str):
        """修改表格单元格"""
        if not 0 <= table_index < len(self._tables):
            raise IndexError(f"表格索引超出范围: {table_index}, 共 {len(self._tables)} 个表格")

        table = self._tables[table_index]
        if not 0 <= row < len(table.rows):
            raise IndexError(f"行索引超出范围: {row}, 表格共 {len(table.rows)} 行")
        if not 0 <= col < len(table.columns):
            raise IndexError(f"列索引超出范围: {col}, 表格共 {len(table.columns)} 列")

        cell = table.cell(row, col)
        # 清空原内容，设置新文本
        cell.text = text

    def _op_replace_table_cell(self, table_index: int, row: int, col: int,
                                pattern: str, replacement: str, regex: bool) -> bool:
        """替换表格单元格文本"""
        if not 0 <= table_index < len(self._tables):
            raise IndexError(f"表格索引超出范围: {table_index}")

        table = self._tables[table_index]
        if not 0 <= row < len(table.rows):
            raise IndexError(f"行索引超出范围: {row}")
        if not 0 <= col < len(table.columns):
            raise IndexError(f"列索引超出范围: {col}")

        cell = table.cell(row, col)
        original = cell.text

        if regex:
            new_text = re.sub(pattern, replacement, original)
        else:
            new_text = original.replace(pattern, replacement)

        if new_text != original:
            cell.text = new_text
            return True
        return False

    def _op_update_table_row(self, table_index: int, row: int, texts: List[str]):
        """批量修改表格整行"""
        if not 0 <= table_index < len(self._tables):
            raise IndexError(f"表格索引超出范围: {table_index}")

        table = self._tables[table_index]
        if not 0 <= row < len(table.rows):
            raise IndexError(f"行索引超出范围: {row}")

        for col, text in enumerate(texts):
            if col < len(table.columns):
                table.cell(row, col).text = text

    def _op_update_table_col(self, table_index: int, col: int, texts: List[str]):
        """批量修改表格整列"""
        if not 0 <= table_index < len(self._tables):
            raise IndexError(f"表格索引超出范围: {table_index}")

        table = self._tables[table_index]
        if not 0 <= col < len(table.columns):
            raise IndexError(f"列索引超出范围: {col}")

        for row, text in enumerate(texts):
            if row < len(table.rows):
                table.cell(row, col).text = text

    # ==================== 图片操作实现 ====================

    def _op_delete_image(self, image_index: int):
        """删除图片"""
        if not 0 <= image_index < len(self._images):
            raise IndexError(f"图片索引超出范围: {image_index}, 共 {len(self._images)} 张图片")

        shape = self._images[image_index]
        shape._inline.getparent().remove(shape._inline)
        self._refresh()

    def _op_resize_image(self, image_index: int, width: Optional[float], height: Optional[float]):
        """调整图片大小（单位：厘米）"""
        if not 0 <= image_index < len(self._images):
            raise IndexError(f"图片索引超出范围: {image_index}, 共 {len(self._images)} 张图片")

        shape = self._images[image_index]

        if width is not None and height is not None:
            # 同时指定宽高
            shape.width = Cm(width)
            shape.height = Cm(height)
        elif width is not None:
            # 只指定宽度，按比例调整高度
            ratio = shape.height / shape.width if shape.width else 1
            shape.width = Cm(width)
            shape.height = int(Cm(width) * ratio)
        elif height is not None:
            # 只指定高度，按比例调整宽度
            ratio = shape.width / shape.height if shape.height else 1
            shape.height = Cm(height)
            shape.width = int(Cm(height) * ratio)

    def _op_insert_image(self, index: int, path: str, width: Optional[float], height: Optional[float]):
        """在段落中插入图片"""
        if not 0 <= index < len(self._paragraphs):
            raise IndexError(f"索引超出范围: {index}")
        if not os.path.exists(path):
            raise FileNotFoundError(f"图片文件不存在: {path}")

        para = self._paragraphs[index]
        run = para.add_run()

        # 插入图片
        if width is not None:
            picture = run.add_picture(path, width=Cm(width))
            if height is not None:
                picture.height = Cm(height)
        elif height is not None:
            picture = run.add_picture(path, height=Cm(height))
        else:
            picture = run.add_picture(path)

        self._refresh()

    # ==================== 引用刷新实现 ====================

    def _op_update_fields_on_open(self):
        """
        设置文档在打开时自动刷新所有域（目录、页码、交叉引用等）

        原理：在 word/settings.xml 中添加 <w:updateFields w:val="true"/>
        这会让 Word 在打开文档时自动计算所有引用
        """
        settings = self.doc.settings.element

        # 检查是否已存在 updateFields 元素
        update_fields = settings.find(qn('w:updateFields'))
        if update_fields is None:
            # 创建新元素
            update_fields = OxmlElement('w:updateFields')
            update_fields.set(qn('w:val'), 'true')
            settings.append(update_fields)
        else:
            # 更新现有元素
            update_fields.set(qn('w:val'), 'true')

    # ==================== 内部辅助 ====================

    def _is_truly_empty(self, para) -> bool:
        """
        判断段落是否真正为空（无文字且无图片等嵌入对象）

        Word 文档中，图片以 <w:drawing> 元素嵌入在段落中，
        这类段落的 para.text 返回空字符串，但删除段落会连带删除图片。

        Returns:
            True: 段落真正为空，可以安全删除
            False: 段落包含嵌入对象，不应作为"空段落"删除
        """
        # 有文字则不是空段落
        if para.text.strip() != "":
            return False

        # 检查是否包含图片（drawing 元素）
        drawings = para._element.findall('.//' + qn('w:drawing'))
        if len(drawings) > 0:
            return False

        # 检查是否包含 OLE 对象（如嵌入的 Excel、公式等）
        objects = para._element.findall('.//' + qn('w:object'))
        if len(objects) > 0:
            return False

        # 检查是否包含图表（Chart 命名空间）
        # 使用通配符匹配，因为图表可能在不同命名空间下
        for child in para._element.iter():
            if 'chart' in child.tag.lower():
                return False

        return True

    def _get_heading_level(self, para) -> Optional[int]:
        """获取标题级别"""
        if not para.style:
            return None
        name = para.style.name
        style_id = para.style.style_id

        if name.startswith('Heading'):
            try:
                return int(name.split()[-1])
            except:
                pass
        if style_id and style_id.isdigit() and int(style_id) <= 9:
            return int(style_id)
        return None

    def _get_section_indices(self, section_title: str) -> List[int]:
        """获取章节内所有段落索引"""
        for i, para in enumerate(self._paragraphs):
            level = self._get_heading_level(para)
            if level and section_title in para.text:
                end = len(self._paragraphs)
                for j in range(i + 1, len(self._paragraphs)):
                    next_level = self._get_heading_level(self._paragraphs[j])
                    if next_level and next_level <= level:
                        end = j
                        break
                return list(range(i, end))
        return []

    def _set_paragraph_text(self, para, text: str):
        """设置段落文本（保留首个run的格式）"""
        if para.runs:
            para.runs[0].text = text
            for r in para.runs[1:]:
                r.text = ''
        else:
            para.add_run(text)

    def _alignment_to_str(self, alignment) -> Optional[str]:
        """对齐方式转字符串"""
        if alignment is None:
            return None
        alignment_map = {
            WD_ALIGN_PARAGRAPH.LEFT: 'left',
            WD_ALIGN_PARAGRAPH.CENTER: 'center',
            WD_ALIGN_PARAGRAPH.RIGHT: 'right',
            WD_ALIGN_PARAGRAPH.JUSTIFY: 'justify',
        }
        return alignment_map.get(alignment)


# ==================== 命令行测试 ====================

def main():
    import sys

    if len(sys.argv) < 2:
        print("用法: python docx_editor.py <file.docx> [outline|read|tables|table|images]")
        return

    editor = DocxEditor(sys.argv[1])
    cmd = sys.argv[2] if len(sys.argv) > 2 else 'outline'

    if cmd == 'outline':
        outline = editor.get_outline()
        print(f"\n文档大纲（共 {outline['total']} 段落）\n" + "=" * 50)
        for h in outline['headings']:
            indent = '  ' * (h['level'] - 1)
            print(f"{indent}[{h['index']:3d}] H{h['level']}: {h['text'][:50]}")

    elif cmd == 'read':
        indices = [int(x) for x in sys.argv[3].split(',')] if len(sys.argv) > 3 else [0, 1, 2]
        content = editor.read_content(indices)
        for p in content:
            print(f"\n[{p['index']}] {p['style']} | empty={p['is_empty']} | numPr={p['xml']['has_numPr']}")
            print(f"  {p['text'][:80]}...")

    elif cmd == 'tables':
        tables = editor.get_tables_outline()
        print(f"\n表格概览（共 {len(tables)} 个表格）\n" + "=" * 50)
        for t in tables:
            print(f"[{t['table_index']}] {t['rows']}行 x {t['cols']}列 | {t['preview']}")

    elif cmd == 'table':
        table_index = int(sys.argv[3]) if len(sys.argv) > 3 else 0
        table = editor.read_table(table_index)
        print(f"\n表格 {table_index}（{table['rows']}行 x {table['cols']}列）\n" + "=" * 50)
        for i, row in enumerate(table['data']):
            print(f"[{i}] {' | '.join(cell[:20] for cell in row)}")

    elif cmd == 'images':
        images = editor.get_images_outline()
        print(f"\n图片概览（共 {len(images)} 张图片）\n" + "=" * 50)
        for img in images:
            print(f"[{img['image_index']}] {img['type']} | {img['width_cm']}cm x {img['height_cm']}cm")


if __name__ == '__main__':
    main()
